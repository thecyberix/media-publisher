from __future__ import annotations

import json
import tempfile
import unittest
from datetime import date, datetime, time
from pathlib import Path
from unittest.mock import MagicMock, patch
from zoneinfo import ZoneInfo

from media_publisher.events.drive_copy import (
    apply_parsed_copy,
    parse_hatha_template_docx,
    parse_program_copy,
)
from media_publisher.events.format import (
    format_bulgarian_datetime,
    format_event_datetime,
    parse_event_date,
    parse_event_time,
)
from media_publisher.events.page import (
    EMPTY_STATE_TEXT,
    SMARTLINK_ACCENT,
    SMARTLINK_BACKGROUND,
    SMARTLINK_TEXT,
    append_event,
    load_events,
    prune_past_events,
    rebuild_index,
)
from media_publisher.events.facebook_event import (
    choose_facebook_image,
    load_image_rotation_state,
    resolve_facebook_image_from_drive,
)
from media_publisher.events.publish import (
    REQUIRED_EVENT_META_SCOPES,
    check_event_meta_scopes,
    publish_event,
)
from media_publisher.events.templates import (
    BHUTA_SHUDDHI_LEARN_MORE_URL,
    EVENT_TYPE_BHUTA_SHUDDHI,
    EVENT_TYPE_SURYA_KRIYA,
    EVENT_TYPE_YOGASANA,
    SURYA_KRIYA_LEARN_MORE_LABEL,
    SURYA_KRIYA_LEARN_MORE_URL,
    YOGASANA_LEARN_MORE_URL,
    city_preposition,
    get_program,
    normalize_event_type,
    render_event,
)
from media_publisher.publishers.meta import MetaClient, MetaError
from media_publisher.sources.google_drive import DriveFile


class EventFormatTests(unittest.TestCase):
    def test_parse_and_format_bulgarian_datetime(self) -> None:
        self.assertEqual(parse_event_date("2026-09-15"), date(2026, 9, 15))
        self.assertEqual(parse_event_time("9:05"), time(9, 5))
        self.assertEqual(
            format_bulgarian_datetime(date(2026, 9, 15), time(18, 0)),
            "15 септември 2026 г., 18:00",
        )
        self.assertEqual(
            format_event_datetime(date(2026, 9, 15), time(18, 0), language="en"),
            "15 September 2026, 18:00",
        )

    def test_rejects_invalid_date_and_time(self) -> None:
        with self.assertRaises(ValueError):
            parse_event_date("15/09/2026")
        with self.assertRaises(ValueError):
            parse_event_time("25:00")


class EventTemplateTests(unittest.TestCase):
    def test_render_surya_kriya_bulgarian(self) -> None:
        rendered = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="София",
            country="България",
            event_date=date(2026, 9, 15),
            event_time=time(18, 0),
            registration_link="https://example.com/register",
        )
        self.assertIn("Суря крия", rendered.title)
        self.assertIn("София", rendered.title)
        self.assertIn("\nв София, България", rendered.title)
        self.assertIn("15 септември 2026 г., 18:00", rendered.full_text)
        self.assertIn("https://example.com/register", rendered.full_text)
        self.assertIn(SURYA_KRIYA_LEARN_MORE_URL, rendered.full_text)
        self.assertIn("https://example.com/register", rendered.facebook_post_text)
        self.assertIn(SURYA_KRIYA_LEARN_MORE_URL, rendered.facebook_post_text)
        self.assertIn("👉 Регистрация тук:", rendered.facebook_post_text)
        self.assertIn(
            f"Вижте какво казва Садгуру: {SURYA_KRIYA_LEARN_MORE_URL}",
            rendered.facebook_post_text,
        )
        self.assertNotIn(SURYA_KRIYA_LEARN_MORE_LABEL, rendered.facebook_post_text)
        self.assertIn(
            '☀️ Програма "Суря крия" в София, България ☀️',
            rendered.facebook_post_text,
        )
        self.assertTrue(
            rendered.facebook_post_text.startswith(
                '☀️ Програма "Суря крия" в София, България ☀️\n'
            )
        )
        self.assertIn("Регистрация", rendered.html_body)
        self.assertIn("<br>в София, България</h2>", rendered.html_body)
        self.assertIn('class="yt-link"', rendered.html_body)
        self.assertIn('class="yt-title"', rendered.html_body)
        self.assertIn(SURYA_KRIYA_LEARN_MORE_LABEL, rendered.html_body)
        self.assertIn(
            f'href="{SURYA_KRIYA_LEARN_MORE_URL}"',
            rendered.html_body,
        )
        self.assertIn("✅ Умствена яснота и фокус", rendered.html_body)
        self.assertIn('class="benefits"', rendered.html_body)

    def test_render_bhuta_shuddhi_bulgarian(self) -> None:
        rendered = render_event(
            event_type=EVENT_TYPE_BHUTA_SHUDDHI,
            city="Пловдив",
            country="България",
            event_date=date(2026, 10, 5),
            event_time=time(11, 0),
            registration_link="https://example.com/bhuta",
        )
        self.assertEqual(rendered.event_type, EVENT_TYPE_BHUTA_SHUDDHI)
        self.assertIn("Бута Шудди", rendered.title)
        self.assertIn("\nв Пловдив, България", rendered.title)
        self.assertIn("петте елемента", rendered.full_text)
        self.assertIn(BHUTA_SHUDDHI_LEARN_MORE_URL, rendered.full_text)
        self.assertIn(
            f"Вижте видеото: {BHUTA_SHUDDHI_LEARN_MORE_URL}",
            rendered.facebook_post_text,
        )
        self.assertIn(
            '💫 Програма "Бута Шудди" в Пловдив, България 💫',
            rendered.facebook_post_text,
        )
        self.assertIn("Регистрация", rendered.html_body)
        self.assertIn("🎯 Хармония и баланс между тялото и ума", rendered.html_body)
        self.assertEqual(rendered.facebook_image_folder, "Bhuta Shuddhi")

    def test_city_preposition_vv_before_v(self) -> None:
        self.assertEqual(city_preposition("София"), "в")
        self.assertEqual(city_preposition("Варна"), "във")
        self.assertEqual(city_preposition("варна"), "във")
        rendered = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="Варна",
            country="България",
            event_date=date(2026, 9, 25),
            event_time=time(10, 0),
            registration_link="https://example.com/varna",
        )
        self.assertIn("\nвъв Варна, България", rendered.title)
        self.assertIn("<br>във Варна, България</h2>", rendered.html_body)

    def test_display_name_event_types(self) -> None:
        self.assertEqual(normalize_event_type("Surya Kriya"), EVENT_TYPE_SURYA_KRIYA)
        self.assertEqual(normalize_event_type("Bhuta Shuddhi"), EVENT_TYPE_BHUTA_SHUDDHI)
        self.assertEqual(normalize_event_type("Yogasanas"), EVENT_TYPE_YOGASANA)
        self.assertEqual(normalize_event_type("yogasana"), EVENT_TYPE_YOGASANA)
        rendered = render_event(
            event_type="Surya Kriya",
            city="София",
            country="България",
            event_date=date(2026, 9, 15),
            event_time=time(18, 0),
            registration_link="https://example.com/register",
        )
        self.assertEqual(rendered.event_type, EVENT_TYPE_SURYA_KRIYA)

    def test_render_yogasana_bulgarian(self) -> None:
        rendered = render_event(
            event_type="Yogasanas",
            city="Пловдив",
            country="България",
            event_date=date(2026, 9, 26),
            event_time=time(9, 0),
            registration_link="https://sadanandayoga.com/events/yogasanas#registration",
        )
        self.assertEqual(rendered.event_type, EVENT_TYPE_YOGASANA)
        self.assertIn("Йогасани", rendered.title)
        self.assertIn("\nв Пловдив, България", rendered.title)
        self.assertIn(YOGASANA_LEARN_MORE_URL, rendered.full_text)
        self.assertIn("✅ Облекчаване на хронични здравословни проблеми", rendered.html_body)
        self.assertEqual(rendered.facebook_image_folder, "Yogasanas")

    def test_unsupported_event_type(self) -> None:
        with self.assertRaises(ValueError):
            render_event(
                event_type="angamardana",
                city="София",
                country="България",
                event_date=date(2026, 9, 15),
                event_time=time(18, 0),
                registration_link="https://example.com/register",
            )


class EventDriveCopyTests(unittest.TestCase):
    def test_parsed_copy_uses_template_quote(self) -> None:
        copy = parse_program_copy(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            english_lines=[
                "☀️ Surya Kriya Programme in [city], [country] ☀️",
                "🗓: [date, month, and time]",
                '"Surya Kriya is a powerful process."- Sadhguru',
                "English body.",
                "Benefits:",
                "✅ Focus",
                "Watch Sadhguru: https://youtu.be/Lh0ZucHjp14",
            ],
            bulgarian_lines=[
                "☀️ Програма “Суря крия” [град], [държава] ☀️",
                "🗓: [дата и час]",
                "„Суря“ означава Слънце, а „крия“ – вътрешен енергиен процес. – Садгуру",
                "Ползи:",
                "✅ Умствена яснота и фокус",
                "Вижте какво казва Садгуру: Заглавие",
            ],
        )
        self.assertEqual(
            copy.quote,
            "„Суря“ означава Слънце, а „крия“ – вътрешен енергиен процес. – Садгуру",
        )
        self.assertEqual(copy.body, "")
        self.assertEqual(copy.benefits, ("Умствена яснота и фокус",))
        program = apply_parsed_copy(get_program(EVENT_TYPE_SURYA_KRIYA), copy)
        rendered = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="София",
            country="България",
            event_date=date(2026, 9, 12),
            event_time=time(8, 0),
            registration_link="https://example.com/register",
            program=program,
        )
        self.assertIn(copy.quote, rendered.full_text)
        self.assertIn(copy.quote, rendered.html_body)
        self.assertNotIn("<p></p>", rendered.html_body)

    def test_joins_sadhguru_attribution_line(self) -> None:
        copy = parse_program_copy(
            event_type=EVENT_TYPE_BHUTA_SHUDDHI,
            english_lines=[
                "💫 Bhuta Shuddhi Programme in [city], [country] 💫",
                "🗓: [date, month, and time]",
                '"English quote."- Sadhguru',
                "English body.",
                "Benefits:",
                "🎯 Harmony",
                "Watch the video: https://youtu.be/jzSX_uBstSA",
            ],
            bulgarian_lines=[
                "💫 Програма “Бута Шудди” в [град], [държава] 💫",
                "🗓: [дата и час]",
                "Бута Шудди е йога система, фокусирана върху пречистването на петте елемента.",
                " – Садгуру",
                "Ползи:",
                "🎯 Хармония и баланс между тялото и ума",
                "Вижте видеото: Заглавие",
            ],
        )
        self.assertEqual(
            copy.quote,
            "Бута Шудди е йога система, фокусирана върху пречистването на петте елемента. – Садгуру",
        )
        self.assertEqual(copy.body, "")

    def test_parse_hatha_template_docx(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "template.docx"
            document = Document()
            english = document.add_table(rows=1, cols=1)
            english.cell(0, 0).text = (
                "☀️ Surya Kriya Programme in [city], [country] ☀️\n"
                "🗓: [date, month, and time]\n"
                '"Surya Kriya is a powerful process."- Sadhguru\n'
                "Surya body.\n"
                "Benefits:\n"
                "✅ Focus\n"
                "Watch Sadhguru: https://youtu.be/Lh0ZucHjp14"
            )
            languages = document.add_table(rows=2, cols=1)
            languages.cell(0, 0).text = "Bulgarian"
            languages.cell(1, 0).text = (
                "☀️ Програма “Суря крия” [град], [държава] ☀️\n"
                "🗓: [дата и час]\n"
                "Цитат на български. – Садгуру\n"
                "Тяло на български.\n"
                "Ползи:\n"
                "✅ Яснота\n"
                "Вижте какво казва Садгуру: Заглавие"
            )
            document.save(path)
            copies = parse_hatha_template_docx(path)
            self.assertIn(EVENT_TYPE_SURYA_KRIYA, copies)
            self.assertEqual(copies[EVENT_TYPE_SURYA_KRIYA].program_name, "Суря крия")
            self.assertEqual(copies[EVENT_TYPE_SURYA_KRIYA].quote, "Цитат на български. – Садгуру")
            self.assertEqual(copies[EVENT_TYPE_SURYA_KRIYA].body, "Тяло на български.")
            self.assertEqual(
                copies[EVENT_TYPE_SURYA_KRIYA].learn_more_url,
                "https://youtu.be/Lh0ZucHjp14",
            )

    def test_parse_hatha_template_docx_uses_bulgarian_hyperlink(self) -> None:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "template.docx"
            document = Document()
            english = document.add_table(rows=1, cols=1)
            english.cell(0, 0).text = (
                "☀️ Surya Kriya Programme in [city], [country] ☀️\n"
                "🗓: [date, month, and time]\n"
                "Benefits:\n"
                "✅ Focus\n"
                "Watch Sadhguru: https://youtu.be/Lh0ZucHjp14"
            )
            languages = document.add_table(rows=2, cols=1)
            languages.cell(0, 0).text = "Bulgarian"
            bg = languages.cell(1, 0)
            bg.text = (
                "☀️ Програма “Суря крия” [град], [държава] ☀️\n"
                "🗓: [дата и час]\n"
                "Цитат. – Садгуру\n"
                "Тяло.\n"
                "Ползи:\n"
                "✅ Яснота\n"
                "Вижте какво казва Садгуру: "
            )
            paragraph = bg.paragraphs[-1]
            rid = paragraph.part.relate_to(
                "https://youtu.be/QFd8S1EHvU8?si=tracking",
                RT.HYPERLINK,
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rid)
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "Суря крия - Запалете Слънцето във вас!"
            run.append(text)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            document.save(path)
            copies = parse_hatha_template_docx(path)
            self.assertEqual(
                copies[EVENT_TYPE_SURYA_KRIYA].learn_more_url,
                "https://youtu.be/QFd8S1EHvU8",
            )

    def test_prefers_bulgarian_youtube_url(self) -> None:
        copy = parse_program_copy(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            english_lines=[
                "☀️ Surya Kriya Programme in [city], [country] ☀️",
                "🗓: [date, month, and time]",
                "Watch Sadhguru: https://youtu.be/Lh0ZucHjp14",
                "Benefits:",
                "✅ Focus",
            ],
            bulgarian_lines=[
                "☀️ Програма “Суря крия” в [град], [държава] ☀️",
                "🗓: [дата и час]",
                "Цитат. – Садгуру",
                "Тяло.",
                "Ползи:",
                "✅ Яснота",
                "Вижте какво казва Садгуру: Заглавие",
            ],
            youtube_url="https://youtu.be/QFd8S1EHvU8?si=tracking",
        )
        self.assertEqual(copy.learn_more_url, "https://youtu.be/QFd8S1EHvU8")


class EventPageTests(unittest.TestCase):
    def test_append_event_and_skip_duplicate(self) -> None:
        rendered = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="Пловдив",
            country="България",
            event_date=date(2026, 10, 1),
            event_time=time(19, 30),
            registration_link="https://example.com/plovdiv",
        )
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            first, created = append_event(
                root,
                rendered,
                facebook_post_id="page_1",
                facebook_permalink="https://www.facebook.com/page_1",
            )
            self.assertTrue(created)
            self.assertTrue((root / "data" / "events.json").is_file())
            self.assertTrue((root / "index.html").is_file())
            html = (root / "index.html").read_text(encoding="utf-8")
            self.assertIn("Пловдив", html)
            self.assertNotIn("Facebook пост", html)
            self.assertIn("Събития", html)
            self.assertIn("assets/sadhguru.png", html)
            self.assertNotIn("Обявени програми", html)
            self.assertNotIn("Доброволци от Иша · Садгуру България", html)
            self.assertIn("justify-content: center", html)
            self.assertIn("border-radius: 16px", html)
            self.assertIn("text-overflow: ellipsis", html)
            self.assertIn("yt-link", html)

            second, created_again = append_event(root, rendered)
            self.assertFalse(created_again)
            self.assertEqual(second.id, first.id)
            events = load_events(root)
            self.assertEqual(len(events), 1)

    def test_append_event_sorts_earliest_first(self) -> None:
        later = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="София",
            country="България",
            event_date=date(2026, 11, 1),
            event_time=time(18, 0),
            registration_link="https://example.com/later",
        )
        earlier = render_event(
            event_type=EVENT_TYPE_BHUTA_SHUDDHI,
            city="Варна",
            country="България",
            event_date=date(2026, 9, 15),
            event_time=time(10, 0),
            registration_link="https://example.com/earlier",
        )
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            append_event(root, later)
            append_event(root, earlier)
            events = load_events(root)
            self.assertEqual(
                [item["city"] for item in events],
                ["Варна", "София"],
            )
            html = (root / "index.html").read_text(encoding="utf-8")
            self.assertLess(html.index("Варна"), html.index("София"))

    def test_rebuild_index_empty(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "data").mkdir(parents=True)
            (root / "data" / "events.json").write_text(
                json.dumps({"events": []}),
                encoding="utf-8",
            )
            path = rebuild_index(root)
            html = path.read_text(encoding="utf-8")
            self.assertIn(EMPTY_STATE_TEXT, html)
            self.assertIn(SMARTLINK_BACKGROUND, html)
            self.assertIn(SMARTLINK_TEXT, html)
            self.assertIn(SMARTLINK_ACCENT, html)
            self.assertIn("Merriweather", html)
            self.assertIn("coming-soon", html)
            self.assertIn("is-empty", html)
            self.assertIn('class="profile"', html)
            self.assertIn("assets/sadhguru.png", html)
            self.assertIn("<h1>Събития</h1>", html)
            self.assertIn("font-size: clamp(1.75rem, 4vw, 2.25rem)", html)
            self.assertEqual(
                html.count("font-size: clamp(1.75rem, 4vw, 2.25rem)"),
                2,
            )

    def test_prune_past_events(self) -> None:
        past = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="София",
            country="България",
            event_date=date(2020, 1, 1),
            event_time=time(10, 0),
            registration_link="https://example.com/past",
        )
        future = render_event(
            event_type=EVENT_TYPE_SURYA_KRIYA,
            city="Варна",
            country="България",
            event_date=date(2099, 6, 1),
            event_time=time(18, 0),
            registration_link="https://example.com/future",
        )
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            before = datetime(2019, 12, 31, 12, 0, tzinfo=ZoneInfo("Europe/Sofia"))
            append_event(root, past, now=before)
            append_event(root, future, now=before)
            self.assertEqual(len(load_events(root)), 2)
            kept, removed = prune_past_events(
                root,
                now=datetime(2026, 8, 11, 12, 0, tzinfo=ZoneInfo("Europe/Sofia")),
            )
            self.assertEqual(len(removed), 1)
            self.assertEqual(removed[0]["city"], "София")
            self.assertEqual(len(kept), 1)
            self.assertEqual(kept[0]["city"], "Варна")
            html = (root / "index.html").read_text(encoding="utf-8")
            self.assertIn("Варна", html)
            self.assertNotIn("София", html)
            self.assertEqual(len(load_events(root)), 1)


class EventPublishTests(unittest.TestCase):
    def test_dry_run_does_not_write(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            project = Path(tmp)
            events_root = project / "events"
            result = publish_event(
                event_type="surya_kriya",
                city="Варна",
                country="България",
                date_text="2026-11-02",
                time_text="17:00",
                registration_link="https://example.com/varna",
                project_root=project,
                events_root=events_root,
                dry_run=True,
            )
            self.assertTrue(result.dry_run)
            self.assertFalse(events_root.exists())

    def test_skip_facebook_writes_page(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            project = Path(tmp)
            events_root = project / "events"
            result = publish_event(
                event_type="surya_kriya",
                city="Варна",
                country="България",
                date_text="2026-11-02",
                time_text="17:00",
                registration_link="https://example.com/varna",
                project_root=project,
                events_root=events_root,
                skip_facebook=True,
            )
            self.assertTrue(result.created_on_page)
            self.assertIsNone(result.facebook_post_id)
            self.assertTrue((events_root / "index.html").is_file())

            again = publish_event(
                event_type="surya_kriya",
                city="Варна",
                country="България",
                date_text="2026-11-02",
                time_text="17:00",
                registration_link="https://example.com/varna",
                project_root=project,
                events_root=events_root,
                skip_facebook=True,
            )
            self.assertTrue(again.skipped_duplicate)
            self.assertFalse(again.created_on_page)

    def test_check_event_meta_scopes_reports_missing(self) -> None:
        fake_info = MagicMock()
        fake_info.scopes = ("pages_manage_posts", "pages_show_list")
        with patch(
            "media_publisher.events.publish.inspect_access_token",
            return_value=fake_info,
        ):
            info, missing = check_event_meta_scopes(
                access_token="token",
                app_id="app",
                app_secret="secret",
                api_version="v21.0",
            )
        self.assertIs(info, fake_info)
        self.assertEqual(missing, [])
        self.assertIn("pages_manage_posts", REQUIRED_EVENT_META_SCOPES)


class EventFacebookImageTests(unittest.TestCase):
    def _surya_images(self) -> list[DriveFile]:
        return [
            DriveFile(id="img1", name="1.jpg", mime_type="image/jpeg"),
            DriveFile(id="img2", name="2.jpg", mime_type="image/jpeg"),
            DriveFile(id="img3", name="3.jpg", mime_type="image/jpeg"),
        ]

    def _drive_with_images(self, images: list[DriveFile] | None = None) -> MagicMock:
        images = images if images is not None else self._surya_images()
        drive = MagicMock()
        drive.find_child_folder.return_value = DriveFile(
            id="folder-surya",
            name="Surya Kriya",
            mime_type="application/vnd.google-apps.folder",
        )
        drive.list_children.return_value = images
        return drive

    def _choose(self, drive, **kwargs):
        kwargs.setdefault("folder_id", "events-root")
        return choose_facebook_image(drive, **kwargs)

    def test_rotation_cycles_through_images(self) -> None:
        drive = self._drive_with_images()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            first, mode1 = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            second, mode2 = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            third, _ = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            fourth, _ = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            self.assertEqual(mode1, "rotation")
            self.assertEqual(mode2, "rotation")
            self.assertEqual([first.id, second.id, third.id], ["img1", "img2", "img3"])
            self.assertEqual(fourth.id, "img1")
            state = load_image_rotation_state(root)
            self.assertEqual(
                state[EVENT_TYPE_SURYA_KRIYA],
                ["img1", "img2", "img3", "img1"],
            )

    def test_explicit_selection_is_skipped_by_later_defaults(self) -> None:
        drive = self._drive_with_images()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            selected, mode = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
                image_id="img3",
            )
            self.assertEqual(mode, "explicit")
            self.assertEqual(selected.id, "img3")

            next_default, next_mode = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            self.assertEqual(next_mode, "rotation")
            self.assertEqual(next_default.id, "img1")

            after_user_gap, _ = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            self.assertEqual(after_user_gap.id, "img2")

            # Cycle complete after img3 (user), img1, img2 — next wraps to img1.
            wrapped, _ = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            self.assertEqual(wrapped.id, "img1")

    def test_explicit_selector_accepts_filename_and_number(self) -> None:
        drive = self._drive_with_images()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            by_name, mode_name = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
                image_id="2.jpg",
            )
            self.assertEqual(mode_name, "explicit")
            self.assertEqual(by_name.id, "img2")

            by_number, mode_number = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
                image_id="1",
            )
            self.assertEqual(mode_number, "explicit")
            self.assertEqual(by_number.id, "img1")

    def test_prior_event_image_ids_seed_rotation(self) -> None:
        drive = self._drive_with_images()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            data_dir = root / "data"
            data_dir.mkdir(parents=True)
            (data_dir / "events.json").write_text(
                json.dumps(
                    {
                        "events": [
                            {
                                "event_type": EVENT_TYPE_SURYA_KRIYA,
                                "facebook_image_id": "img2",
                                "created_at": "2026-01-01T10:00:00+00:00",
                            },
                            {
                                "event_type": EVENT_TYPE_SURYA_KRIYA,
                                "facebook_image_id": "img1",
                                "created_at": "2026-01-02T10:00:00+00:00",
                            },
                        ]
                    }
                ),
                encoding="utf-8",
            )
            selected, mode = self._choose(
                drive,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                events_root=root,
            )
            self.assertEqual(mode, "rotation")
            # img1 and img2 already used by prior events → default is img3.
            self.assertEqual(selected.id, "img3")

    def test_explicit_image_id_must_be_in_folder(self) -> None:
        drive = self._drive_with_images()
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            with self.assertRaisesRegex(Exception, "not in the"):
                self._choose(
                    drive,
                    event_type=EVENT_TYPE_SURYA_KRIYA,
                    events_root=root,
                    image_id="missing",
                )

    def test_resolve_facebook_image_from_drive(self) -> None:
        drive = self._drive_with_images()

        def _download(file_id: str, destination: Path) -> Path:
            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(b"fake-image")
            return destination

        drive.download_file.side_effect = _download
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            events_root = root / "events"
            events_root.mkdir()
            selected = resolve_facebook_image_from_drive(
                project_root=root,
                events_root=events_root,
                event_type=EVENT_TYPE_SURYA_KRIYA,
                drive_client=drive,
                folder_id="events-root",
            )
            self.assertTrue(selected.local_path.is_file())
            self.assertEqual(selected.drive_file.id, "img1")
            self.assertEqual(selected.selection, "rotation")
            self.assertIn("img1", selected.local_path.name)
            drive.find_child_folder.assert_called_once()
            drive.download_file.assert_called_once()


class MetaFeedAndCommentTests(unittest.TestCase):
    def test_create_facebook_feed_post(self) -> None:
        client = MetaClient("token-test", app_id="app123")
        with patch.object(client, "_request", return_value={"id": "page_99"}) as request_mock:
            post_id = client.create_facebook_feed_post(
                page_id="page123",
                message="Hello event",
            )
        self.assertEqual(post_id, "page_99")
        request_mock.assert_called_once_with(
            "POST",
            "page123/feed",
            body={"message": "Hello event", "published": "true"},
        )

    def test_create_facebook_photo_post(self) -> None:
        client = MetaClient("token-test", app_id="app123")
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "event.jpg"
            image.write_bytes(b"fake-image")
            with patch.object(
                client,
                "_multipart_request",
                return_value={"id": "photo_1", "post_id": "page123_99"},
            ) as request_mock:
                post_id = client.create_facebook_photo_post(
                    page_id="page123",
                    message="Hello event",
                    image_path=image,
                )
        self.assertEqual(post_id, "page123_99")
        request_mock.assert_called_once()
        args, kwargs = request_mock.call_args
        self.assertEqual(args[0], "page123/photos")
        self.assertEqual(kwargs["fields"]["published"], "true")
        self.assertEqual(kwargs["fields"]["caption"], "Hello event")
        self.assertIn("source", kwargs["files"])

    def test_create_facebook_comment(self) -> None:
        client = MetaClient("token-test", app_id="app123")
        with patch.object(client, "_request", return_value={"id": "comment_1"}) as request_mock:
            comment_id = client.create_facebook_comment(
                object_id="page_99",
                message="👉 Регистрация тук: https://example.com",
            )
        self.assertEqual(comment_id, "comment_1")
        request_mock.assert_called_once_with(
            "POST",
            "page_99/comments",
            body={"message": "👉 Регистрация тук: https://example.com"},
        )

    def test_create_facebook_feed_post_requires_message(self) -> None:
        client = MetaClient("token-test", app_id="app123")
        with self.assertRaises(MetaError):
            client.create_facebook_feed_post(page_id="page123", message="  ")


if __name__ == "__main__":
    unittest.main()
