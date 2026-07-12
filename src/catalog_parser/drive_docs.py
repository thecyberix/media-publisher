from __future__ import annotations

import io
import re
from typing import Any, Iterator

from googleapiclient.discovery import Resource

from catalog_parser.canva import extract_canva_design_url

GOOGLE_DOC_MIME_TYPE = "application/vnd.google-apps.document"
WORD_DOC_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_TEXT_MIME_TYPES = (GOOGLE_DOC_MIME_TYPE, WORD_DOC_MIME_TYPE)
DRIVE_FOLDER_PATTERN = re.compile(r"/folders/([a-zA-Z0-9_-]+)")
DRIVE_FILE_PATTERN = re.compile(r"/file/d/([a-zA-Z0-9_-]+)")

TITLE_YT_TABLE_LABEL = "TITLE - YT"
DESCRIPTION_TABLE_LABEL = "Description"
THUMBNAIL_TABLE_LABEL = "THUMBNAIL - YT"
DEFAULT_YT_TITLE_FIELD = "ytTitle"
DEFAULT_YT_DESCRIPTION_FIELD = "ytDescription"
DEFAULT_YT_THUMBNAIL_FIELD = "ytThumbnail"

LABELED_TABLE_FIELDS = {
    TITLE_YT_TABLE_LABEL: DEFAULT_YT_TITLE_FIELD,
    DESCRIPTION_TABLE_LABEL: DEFAULT_YT_DESCRIPTION_FIELD,
    THUMBNAIL_TABLE_LABEL: DEFAULT_YT_THUMBNAIL_FIELD,
}
THUMBNAIL_TABLE_LABEL_ALIASES = (
    THUMBNAIL_TABLE_LABEL,
    "Thumbnail",
    "THUMB - YT",
)


class DriveDocsError(RuntimeError):
    pass


def extract_drive_folder_id(value: str) -> str | None:
    value = value.strip()
    if not value:
        return None

    match = DRIVE_FOLDER_PATTERN.search(value)
    if match:
        return match.group(1)

    if re.fullmatch(r"[a-zA-Z0-9_-]+", value):
        return value
    return None


def extract_drive_file_id(value: str) -> str | None:
    value = value.strip()
    if not value:
        return None

    match = DRIVE_FILE_PATTERN.search(value)
    if match:
        return match.group(1)

    if re.fullmatch(r"[a-zA-Z0-9_-]+", value):
        return value
    return None


def drive_file_view_url(file_id: str) -> str:
    return f"https://drive.google.com/file/d/{file_id}/view"


def drive_file_download_url(file_id: str) -> str:
    return f"https://drive.google.com/uc?export=download&id={file_id}"


def _normalize_label(value: str) -> str:
    return " ".join(value.strip().casefold().split())


def _paragraph_text(element: dict[str, Any]) -> str:
    paragraph = element.get("paragraph")
    if not isinstance(paragraph, dict):
        return ""

    parts: list[str] = []
    for item in paragraph.get("elements", []):
        if not isinstance(item, dict):
            continue
        text_run = item.get("textRun")
        if isinstance(text_run, dict):
            parts.append(str(text_run.get("content", "")))
    return "".join(parts).strip()


def table_to_grid(table: dict[str, Any]) -> list[list[str]]:
    grid: list[list[str]] = []
    for row in table.get("tableRows", []):
        if not isinstance(row, dict):
            continue
        cells: list[str] = []
        for cell in row.get("tableCells", []):
            if not isinstance(cell, dict):
                cells.append("")
                continue
            cell_parts: list[str] = []
            for content in cell.get("content", []):
                if isinstance(content, dict):
                    text = _paragraph_text(content)
                    if text:
                        cell_parts.append(text)
            cells.append("\n".join(cell_parts).strip())
        grid.append(cells)
    return grid


def docx_table_to_grid(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def table_matches_label(
    grid: list[list[str]],
    preceding_text: str | None,
    label: str,
) -> bool:
    if preceding_text and _normalize_label(preceding_text) == _normalize_label(label):
        return True

    if not grid:
        return False

    first_row = " ".join(grid[0]).strip()
    return _normalize_label(label) in _normalize_label(first_row)


def table_matches_title_yt(grid: list[list[str]], preceding_text: str | None) -> bool:
    return table_matches_label(grid, preceding_text, TITLE_YT_TABLE_LABEL)


def extract_table_value_from_grid(grid: list[list[str]]) -> str | None:
    if len(grid) < 2 or not grid[1]:
        return None

    value = grid[1][0].strip()
    return value or None


def extract_yt_title_from_grid(grid: list[list[str]]) -> str | None:
    return extract_table_value_from_grid(grid)


def extract_labeled_fields_from_blocks(
    blocks: Iterator[tuple[str, Any]],
    *,
    label_fields: dict[str, str] | None = None,
) -> dict[str, str | None]:
    mappings = label_fields or LABELED_TABLE_FIELDS
    results = {field_name: None for field_name in mappings.values()}
    previous_text = ""

    for block_type, block in blocks:
        if block_type == "paragraph":
            text = str(block).strip()
            if text:
                previous_text = text
            continue

        grid = block
        matched_field: str | None = None
        for label, field_name in mappings.items():
            if table_matches_label(grid, previous_text, label):
                matched_field = field_name
                break
        if matched_field is None:
            for alias in THUMBNAIL_TABLE_LABEL_ALIASES:
                if table_matches_label(grid, previous_text, alias):
                    matched_field = DEFAULT_YT_THUMBNAIL_FIELD
                    break

        previous_text = ""
        if matched_field is None:
            continue

        value = extract_table_value_from_grid(grid)
        if value and results.get(matched_field) is None:
            results[matched_field] = value

    return results


def extract_drive_fields_from_document(document: dict[str, Any]) -> dict[str, str | None]:
    content = document.get("body", {}).get("content", [])
    if not isinstance(content, list):
        return {field: None for field in LABELED_TABLE_FIELDS.values()}

    def iter_google_blocks() -> Iterator[tuple[str, Any]]:
        for element in content:
            if not isinstance(element, dict):
                continue
            if "paragraph" in element:
                yield "paragraph", _paragraph_text(element)
                continue
            if "table" not in element:
                continue
            table = element["table"]
            if isinstance(table, dict):
                yield "table", table_to_grid(table)

    return extract_labeled_fields_from_blocks(iter_google_blocks())


def extract_yt_title_from_document(document: dict[str, Any]) -> str | None:
    return extract_drive_fields_from_document(document).get(DEFAULT_YT_TITLE_FIELD)


def _iter_docx_blocks(document: Any) -> Iterator[tuple[str, Any]]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document).text
        elif isinstance(child, CT_Tbl):
            yield "table", docx_table_to_grid(Table(child, document))


def extract_drive_fields_from_docx(document: Any) -> dict[str, str | None]:
    fields = extract_labeled_fields_from_blocks(_iter_docx_blocks(document))
    hyperlink = _extract_first_canva_hyperlink_from_docx(document)
    if hyperlink and not fields.get(DEFAULT_YT_THUMBNAIL_FIELD):
        fields[DEFAULT_YT_THUMBNAIL_FIELD] = hyperlink
    return fields


def _extract_first_canva_hyperlink_from_docx(document: Any) -> str | None:
    from docx.oxml.ns import qn

    for paragraph in document.paragraphs:
        hyperlink = _extract_canva_hyperlink_from_paragraph(paragraph, qn)
        if hyperlink:
            return hyperlink

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    hyperlink = _extract_canva_hyperlink_from_paragraph(paragraph, qn)
                    if hyperlink:
                        return hyperlink
    return None


def _extract_canva_hyperlink_from_paragraph(paragraph: Any, qn: Any) -> str | None:
    for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
        relationship_id = hyperlink.get(qn("r:id"))
        if not relationship_id:
            continue
        relationship = paragraph.part.rels.get(relationship_id)
        if relationship is None:
            continue
        target = getattr(relationship, "target_ref", None)
        if not isinstance(target, str):
            continue
        canva_url = extract_canva_design_url(target)
        if canva_url:
            return canva_url
    return None


def extract_yt_title_from_docx(document: Any) -> str | None:
    return extract_drive_fields_from_docx(document).get(DEFAULT_YT_TITLE_FIELD)


def list_text_documents_in_folder(
    drive_service: Resource,
    folder_id: str,
) -> list[dict[str, str]]:
    mime_query = " or ".join(f"mimeType='{mime_type}'" for mime_type in SUPPORTED_TEXT_MIME_TYPES)
    query = f"'{folder_id}' in parents and ({mime_query}) and trashed=false"
    response = (
        drive_service.files()
        .list(
            q=query,
            fields="files(id,name,mimeType)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        )
        .execute()
    )
    files = response.get("files", [])
    if not isinstance(files, list):
        return []

    documents: list[dict[str, str]] = []
    for file_info in files:
        if not isinstance(file_info, dict):
            continue
        file_id = file_info.get("id")
        file_name = file_info.get("name")
        mime_type = file_info.get("mimeType")
        if isinstance(file_id, str) and file_id and isinstance(mime_type, str):
            documents.append(
                {
                    "id": file_id,
                    "name": file_name if isinstance(file_name, str) else "",
                    "mimeType": mime_type,
                }
            )

    documents.sort(key=_document_sort_key)
    return documents


def _document_sort_key(document: dict[str, str]) -> tuple[int, int, str]:
    name = document.get("name", "")
    mime_type = document.get("mimeType", "")
    text_prefix_rank = 0 if name.upper().startswith("TEXT_") else 1
    mime_rank = 0 if mime_type == GOOGLE_DOC_MIME_TYPE else 1
    return (text_prefix_rank, mime_rank, name.casefold())


def read_drive_fields_from_google_doc(
    docs_service: Resource,
    document_id: str,
) -> dict[str, str | None]:
    document = docs_service.documents().get(documentId=document_id).execute()
    if not isinstance(document, dict):
        raise DriveDocsError(f"Unexpected Docs API response for document {document_id!r}")
    return extract_drive_fields_from_document(document)


def read_drive_fields_from_word_doc(
    drive_service: Resource,
    document_id: str,
) -> dict[str, str | None]:
    from docx import Document

    content = drive_service.files().get_media(fileId=document_id).execute()
    document = Document(io.BytesIO(content))
    return extract_drive_fields_from_docx(document)


def read_drive_fields_from_file(
    drive_service: Resource,
    docs_service: Resource | None,
    document: dict[str, str],
) -> dict[str, str | None]:
    mime_type = document.get("mimeType")
    document_id = document.get("id")
    if not document_id:
        raise DriveDocsError(f"Document {document.get('name')!r} is missing an id")

    if mime_type == GOOGLE_DOC_MIME_TYPE:
        if docs_service is None:
            raise DriveDocsError("Google Docs API client is not configured")
        return read_drive_fields_from_google_doc(docs_service, document_id)

    if mime_type == WORD_DOC_MIME_TYPE:
        return read_drive_fields_from_word_doc(drive_service, document_id)

    raise DriveDocsError(f"Unsupported document type: {mime_type!r}")


def _fields_have_values(fields: dict[str, str | None]) -> bool:
    return any(value for value in fields.values())


def read_drive_fields_from_folder(
    drive_service: Resource,
    docs_service: Resource | None,
    folder_id: str,
) -> dict[str, str | None]:
    documents = list_text_documents_in_folder(drive_service, folder_id)
    if not documents:
        raise DriveDocsError(
            f"No Google Docs or Word documents found in Drive folder {folder_id!r}"
        )

    last_error: str | None = None
    merged = {field: None for field in LABELED_TABLE_FIELDS.values()}

    for document in documents:
        try:
            fields = read_drive_fields_from_file(drive_service, docs_service, document)
        except Exception as exc:
            last_error = str(exc)
            continue

        for field_name, value in fields.items():
            if value and merged.get(field_name) is None:
                merged[field_name] = value

        if _fields_have_values(merged):
            return merged

        last_error = (
            f"Document {document['name']!r} has no labeled table values"
        )

    if last_error and not _fields_have_values(merged):
        raise DriveDocsError(last_error)
    return merged


def read_yt_title_from_folder(
    drive_service: Resource,
    docs_service: Resource | None,
    folder_id: str,
) -> str | None:
    return read_drive_fields_from_folder(drive_service, docs_service, folder_id).get(
        DEFAULT_YT_TITLE_FIELD
    )


def enrich_records_with_yt_titles(
    records: list[dict[str, Any]],
    drive_service: Resource,
    docs_service: Resource | None,
    *,
    folder_link_field: str = "pkgLink",
    title_field: str = DEFAULT_YT_TITLE_FIELD,
    description_field: str = DEFAULT_YT_DESCRIPTION_FIELD,
) -> list[dict[str, Any]]:
    enriched: list[dict[str, Any]] = []
    total = len(records)

    for index, record in enumerate(records, start=1):
        updated = dict(record)
        folder_link = updated.get(folder_link_field)
        label = updated.get("ctTitle")
        label_text = label if isinstance(label, str) and label else f"row {index}"
        print(f"Drive {index}/{total}: {label_text}")

        if not isinstance(folder_link, str) or not folder_link.strip():
            updated[title_field] = None
            updated[description_field] = None
            updated[f"{title_field}Error"] = f"Missing {folder_link_field}"
            enriched.append(updated)
            continue

        folder_id = extract_drive_folder_id(folder_link)
        if folder_id is None:
            updated[title_field] = None
            updated[description_field] = None
            updated[f"{title_field}Error"] = (
                f"Could not parse Drive folder id from {folder_link!r}"
            )
            enriched.append(updated)
            continue

        try:
            fields = read_drive_fields_from_folder(
                drive_service,
                docs_service,
                folder_id,
            )
            updated[title_field] = fields.get(title_field)
            updated[description_field] = fields.get(description_field)
            updated.pop(f"{title_field}Error", None)
            if updated[title_field]:
                print(f"  -> {title_field}: {updated[title_field]}")
            if updated[description_field]:
                print(f"  -> {description_field}: {updated[description_field]}")
            if not updated[title_field] and not updated[description_field]:
                print("  -> not found")
        except DriveDocsError as exc:
            updated[title_field] = None
            updated[description_field] = None
            updated[f"{title_field}Error"] = str(exc)
            print(f"  -> error: {exc}")
        except Exception as exc:
            updated[title_field] = None
            updated[description_field] = None
            updated[f"{title_field}Error"] = str(exc)
            print(f"  -> error: {exc}")

        enriched.append(updated)

    return enriched
