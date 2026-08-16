from __future__ import annotations

import re
from dataclasses import dataclass, replace
from pathlib import Path

from media_publisher.events.templates import (
    EVENT_TYPE_BHUTA_SHUDDHI,
    EVENT_TYPE_SURYA_KRIYA,
    EVENT_TYPE_YOGASANA,
    PROGRAMS,
    ProgramTemplate,
    get_program,
    normalize_event_type,
)
from media_publisher.languages import LanguageDefinition, selected_language
from media_publisher.sources.drive_layout import resolve_events_folder_id
from media_publisher.sources.google_drive import (
    DOCX_MIME_TYPE,
    FOLDER_MIME_TYPE,
    GOOGLE_DOC_MIME_TYPE,
    DriveFile,
    GoogleDriveClient,
    GoogleDriveError,
)

TEMPLATE_CACHE_NAME = "hatha-message-template.docx"
ENGLISH_HEADING_RE = re.compile(
    r"(?P<name>Surya Kriya|Bhuta Shuddhi|Yogasanas?|Angamardana)\s+Programme",
    re.IGNORECASE,
)
LEARN_MORE_URL_RE = re.compile(r"https?://\S+")
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
ENGLISH_BENEFITS_HEADINGS = ("benefits", "this programme offers")
ENGLISH_LEARN_MORE_PREFIXES = ("watch", "learn more", "find out")

EVENT_TYPE_BY_ENGLISH_NAME = {
    "surya kriya": EVENT_TYPE_SURYA_KRIYA,
    "bhuta shuddhi": EVENT_TYPE_BHUTA_SHUDDHI,
    "yogasana": EVENT_TYPE_YOGASANA,
    "yogasanas": EVENT_TYPE_YOGASANA,
    "angamardana": "angamardana",
}


class EventTemplateError(RuntimeError):
    pass


@dataclass(frozen=True)
class ParsedProgramCopy:
    event_type: str
    program_name: str
    title_emoji: str
    quote: str
    body: str
    benefits: tuple[str, ...]
    benefit_bullet: str
    learn_more_intro: str
    learn_more_url: str
    learn_more_label: str
    benefits_heading: str


def event_template_cache_path(project_root: Path) -> Path:
    return project_root / "downloads" / "event-images" / TEMPLATE_CACHE_NAME


def find_event_template_file(
    drive_client: GoogleDriveClient,
    folder_id: str = "",
) -> DriveFile:
    if not folder_id.strip():
        folder_id = resolve_events_folder_id(drive_client)
    try:
        children = drive_client.list_children(folder_id)
    except GoogleDriveError as exc:
        raise EventTemplateError(
            f"Failed to list event templates in Drive folder {folder_id}: {exc}"
        ) from exc
    docs = [
        item
        for item in children
        if item.mime_type != FOLDER_MIME_TYPE
        and (
            item.mime_type in {GOOGLE_DOC_MIME_TYPE, DOCX_MIME_TYPE}
            or item.name.casefold().endswith(".docx")
        )
    ]
    if not docs:
        raise EventTemplateError(
            f"No Hatha message template (.docx / Google Doc) found in Drive folder {folder_id}"
        )

    def _rank(item: DriveFile) -> tuple[int, int, int, str]:
        name = item.name.casefold()
        return (
            0 if "hatha" in name else 1,
            0 if "template" in name else 1,
            0 if item.mime_type == GOOGLE_DOC_MIME_TYPE else 1,
            name,
        )

    return sorted(docs, key=_rank)[0]


def load_program_from_drive(
    drive_client: GoogleDriveClient,
    event_type: str,
    *,
    project_root: Path,
    folder_id: str = "",
) -> ProgramTemplate:
    base = get_program(event_type)
    copies = load_program_copies_from_drive(
        drive_client,
        project_root=project_root,
        folder_id=folder_id,
    )
    copy = copies.get(base.event_type)
    if copy is None:
        language = selected_language()
        raise EventTemplateError(
            f"Drive template has no {language.name} copy for {base.event_type!r}"
        )
    return apply_parsed_copy(base, copy)


def load_program_copies_from_drive(
    drive_client: GoogleDriveClient,
    *,
    project_root: Path,
    folder_id: str = "",
) -> dict[str, ParsedProgramCopy]:
    template = find_event_template_file(drive_client, folder_id)
    destination = event_template_cache_path(project_root)
    try:
        drive_client.download_document(template.id, destination)
    except GoogleDriveError as exc:
        raise EventTemplateError(
            f"Failed to download event template {template.name!r}: {exc}"
        ) from exc
    return parse_hatha_template_docx(destination)


def parse_hatha_template_docx(
    path: Path,
    *,
    language: LanguageDefinition | None = None,
) -> dict[str, ParsedProgramCopy]:
    try:
        from docx import Document
    except ImportError as exc:
        raise EventTemplateError(
            "Parsing the Hatha template requires python-docx"
        ) from exc
    document = Document(str(path))
    copies: dict[str, ParsedProgramCopy] = {}
    tables = list(document.tables)
    definition = language or selected_language()
    for index, table in enumerate(tables):
        if len(table.rows) != 1 or len(table.columns) != 1:
            continue
        english_cell = table.rows[0].cells[0]
        english_lines = _cell_lines(english_cell)
        heading = next((line for line in english_lines if line.strip()), "")
        event_type = event_type_from_english_heading(heading)
        if event_type is None or event_type not in PROGRAMS:
            continue
        language_cell = _find_language_cell(
            tables[index + 1 : index + 4],
            language=definition,
        )
        if language_cell is None:
            continue
        copies[event_type] = parse_program_copy(
            event_type=event_type,
            english_lines=english_lines,
            language_lines=_cell_lines(language_cell),
            youtube_url=(
                _cell_youtube_url(language_cell) or _cell_youtube_url(english_cell)
            ),
            language=definition,
        )
    if not copies:
        raise EventTemplateError("Hatha template did not contain any programme copy")
    return copies


def event_type_from_english_heading(heading: str) -> str | None:
    match = ENGLISH_HEADING_RE.search(heading)
    if match is None:
        return None
    return EVENT_TYPE_BY_ENGLISH_NAME.get(match.group("name").casefold())


def parse_program_copy(
    *,
    event_type: str,
    english_lines: list[str],
    language_lines: list[str],
    youtube_url: str | None = None,
    language: LanguageDefinition | None = None,
) -> ParsedProgramCopy:
    definition = language or selected_language()
    events = definition.require_events()
    english = _parse_language_block(
        english_lines,
        benefits_headings=ENGLISH_BENEFITS_HEADINGS,
        learn_more_prefixes=ENGLISH_LEARN_MORE_PREFIXES,
        quote_attributions=(),
    )
    localized = _parse_language_block(
        language_lines,
        benefits_headings=events.benefits_headings,
        learn_more_prefixes=events.learn_more_prefixes,
        quote_attributions=events.quote_attributions,
    )
    title = next((line for line in language_lines if line.strip()), "")
    program_name = _program_name_from_title(title, events.program_word)
    if not program_name:
        raise EventTemplateError(
            f"Could not parse {definition.name} programme name for {event_type!r}"
        )
    quote = localized.quote
    learn_more_url = _normalize_youtube_url(
        (youtube_url or "").strip()
        or localized.learn_more_url
        or english.learn_more_url
    )
    if not learn_more_url:
        raise EventTemplateError(
            f"Template is missing a YouTube URL for {event_type!r}"
        )
    if not quote and not localized.body:
        raise EventTemplateError(
            f"{definition.name} template is missing quote/body for {event_type!r}"
        )
    if not localized.benefits:
        raise EventTemplateError(
            f"{definition.name} template is missing benefits for {event_type!r}"
        )
    return ParsedProgramCopy(
        event_type=normalize_event_type(event_type),
        program_name=program_name,
        title_emoji=_title_emoji(title, events.program_word) or _title_emoji(
            next((line for line in english_lines if line.strip()), ""),
            "Programme",
        ),
        quote=quote,
        body=localized.body,
        benefits=localized.benefits,
        benefit_bullet=localized.benefit_bullet,
        learn_more_intro=localized.learn_more_intro,
        learn_more_url=learn_more_url,
        learn_more_label=localized.learn_more_label,
        benefits_heading=localized.benefits_heading,
    )


def apply_parsed_copy(base: ProgramTemplate, copy: ParsedProgramCopy) -> ProgramTemplate:
    return replace(
        base,
        program_name=copy.program_name,
        title_emoji=copy.title_emoji or base.title_emoji,
        quote=copy.quote,
        body=copy.body,
        benefits=copy.benefits,
        benefit_bullet=copy.benefit_bullet or base.benefit_bullet,
        learn_more_intro=copy.learn_more_intro,
        learn_more_url=copy.learn_more_url,
        learn_more_label=copy.learn_more_label,
        benefits_heading=copy.benefits_heading or base.benefits_heading,
    )


@dataclass(frozen=True)
class _LanguageBlock:
    quote: str
    body: str
    benefits: tuple[str, ...]
    benefit_bullet: str
    learn_more_intro: str
    learn_more_url: str
    learn_more_label: str
    benefits_heading: str


def _parse_language_block(
    lines: list[str],
    *,
    benefits_headings: tuple[str, ...],
    learn_more_prefixes: tuple[str, ...],
    quote_attributions: tuple[str, ...],
) -> _LanguageBlock:
    stripped = [line.strip() for line in lines]
    quote = ""
    body = ""
    benefits: list[str] = []
    benefit_bullet = ""
    learn_more_intro = ""
    learn_more_url = ""
    learn_more_label = ""
    benefits_heading = ""
    mode = "pre"
    heading_set = {heading.casefold().rstrip(":") for heading in benefits_headings}
    attribution_re = _attribution_re(quote_attributions)
    for line in stripped:
        if not line:
            continue
        if line.startswith("🗓") or line.startswith("🗓️"):
            mode = "quote"
            continue
        if line.casefold().rstrip(":") in heading_set:
            benefits_heading = line
            mode = "benefits"
            continue
        if _is_learn_more_line(line, learn_more_prefixes):
            intro, url, label = _split_learn_more(line)
            learn_more_intro = intro
            learn_more_url = url
            learn_more_label = label
            mode = "done"
            continue
        if line.startswith("👉") or line.startswith("💫"):
            mode = "done"
            continue
        if mode == "quote":
            quote = _clean_quote(line)
            mode = "quote_tail"
            continue
        if (
            mode == "quote_tail"
            and attribution_re is not None
            and attribution_re.match(line)
        ):
            quote = f"{quote} {line}".strip()
            mode = "body"
            continue
        if mode in {"quote_tail", "body"} and not body:
            body = line
            mode = "body"
            continue
        if mode == "benefits":
            bullet, text = _split_benefit(line)
            if text:
                if bullet and not benefit_bullet:
                    benefit_bullet = bullet
                benefits.append(text)
    return _LanguageBlock(
        quote=quote,
        body=body,
        benefits=tuple(benefits),
        benefit_bullet=benefit_bullet,
        learn_more_intro=learn_more_intro,
        learn_more_url=learn_more_url,
        learn_more_label=learn_more_label,
        benefits_heading=benefits_heading,
    )


def _attribution_re(names: tuple[str, ...]) -> re.Pattern[str] | None:
    cleaned = [re.escape(name.strip()) for name in names if name.strip()]
    if not cleaned:
        return None
    return re.compile(
        rf"^[-–—]\s*(?:{'|'.join(cleaned)})\s*$",
        re.IGNORECASE,
    )


def _find_language_cell(tables, *, language: LanguageDefinition):
    header_name = language.name.casefold()
    for table in tables:
        if not table.rows:
            continue
        header = table.rows[0].cells[0].text.strip().casefold()
        if header != header_name or len(table.rows) < 2:
            continue
        return table.rows[1].cells[0]
    return None


def _cell_youtube_url(cell) -> str:
    from docx.oxml.ns import qn

    rels = getattr(getattr(cell, "part", None), "rels", None)
    if rels is None:
        return ""
    getter = getattr(rels, "get", None)
    for node in cell._tc.iter():
        if not str(node.tag).endswith("}hyperlink"):
            continue
        rid = node.get(qn("r:id"))
        if not rid:
            continue
        relationship = getter(rid) if callable(getter) else None
        if relationship is None:
            continue
        target = str(getattr(relationship, "target_ref", "") or "")
        if "youtu" in target.casefold():
            return _normalize_youtube_url(target)
    return ""


def _normalize_youtube_url(url: str) -> str:
    cleaned = url.strip()
    for separator in ("?", "#"):
        if separator in cleaned:
            cleaned = cleaned.split(separator, 1)[0]
    return cleaned.rstrip("/")


def _cell_lines(cell) -> list[str]:
    lines: list[str] = []
    for paragraph in cell.paragraphs:
        lines.extend(_paragraph_lines(paragraph))
    return lines


def _paragraph_lines(paragraph) -> list[str]:
    parts: list[str] = [""]
    for node in paragraph._element.iter():
        if node.tag in {f"{W_NS}br", f"{W_NS}cr"}:
            parts.append("")
        elif node.tag == f"{W_NS}t" and node.text:
            parts[-1] += node.text
    if len(parts) == 1 and "\n" in parts[0]:
        return parts[0].splitlines()
    return parts


def _program_name_from_title(title: str, program_word: str) -> str:
    pattern = re.compile(
        re.escape(program_word) + r'\s+[“"„«]\s*(?P<name>[^”"»“]+?)\s*[”"»“]',
        re.IGNORECASE,
    )
    match = pattern.search(title)
    if match is None:
        return ""
    return match.group("name").strip()


def _title_emoji(title: str, program_word: str) -> str:
    pattern = re.compile(r"^(\S+)\s+" + re.escape(program_word), re.IGNORECASE)
    match = pattern.match(title.strip())
    if match:
        return match.group(1)
    match = re.match(r"^(\S+)\s+", title.strip())
    return match.group(1) if match else ""


def _clean_quote(text: str) -> str:
    return text.strip().strip("_").strip()


def _is_learn_more_line(line: str, prefixes: tuple[str, ...]) -> bool:
    folded = line.casefold()
    if LEARN_MORE_URL_RE.search(line):
        return True
    return any(folded.startswith(prefix.casefold()) for prefix in prefixes)


def _split_learn_more(line: str) -> tuple[str, str, str]:
    url_match = LEARN_MORE_URL_RE.search(line)
    url = url_match.group(0).rstrip(".,)") if url_match else ""
    remainder = line
    if url:
        remainder = (line[: url_match.start()] + line[url_match.end() :]).strip()
    intro, _, label = remainder.partition(":")
    intro = intro.strip()
    if intro and not intro.endswith(":"):
        intro = f"{intro}:"
    return intro, url, label.strip()


def _split_benefit(line: str) -> tuple[str, str]:
    parts = line.strip().split(None, 1)
    if len(parts) == 2 and not parts[0][:1].isalnum():
        return parts[0], parts[1]
    return "", line.strip()
