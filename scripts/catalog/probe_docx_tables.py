from __future__ import annotations

import io
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from docx import Document

from catalog_parser.auth import get_drive_service
from catalog_parser.__main__ import DEFAULT_CREDENTIALS, DEFAULT_TOKEN, load_env_file

load_env_file(PROJECT_ROOT / ".env")

file_id = "1gHFU73i2FdP29oQRcD07kUc7Ns0gm2Vb"
drive = get_drive_service(DEFAULT_CREDENTIALS, DEFAULT_TOKEN)
content = drive.files().get_media(fileId=file_id).execute()
document = Document(io.BytesIO(content))
for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index}: {text[:120]}")
for index, table in enumerate(document.tables):
    print(f"TABLE {index}")
    for row in table.rows:
        print([cell.text.strip() for cell in row.cells])
