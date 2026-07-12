"""Download pkgTn-marked unpublished thumbnails into tn-cache as Airtable-ready JPGs."""
from __future__ import annotations

import io
import json
import re
import sys
import tempfile
from collections import Counter, defaultdict
from pathlib import Path
from urllib.parse import urlparse, urlunparse

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))
sys.path.insert(0, str(PROJECT_ROOT / "src/catalog_parser"))

from catalog_parser.parser import rows_to_records
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from google.oauth2 import service_account
from googleapiclient.discovery import build
from PIL import Image

from media_publisher.__main__ import canva_client_from_settings, canva_settings_complete
from media_publisher.config import load_settings
from media_publisher.sources.airtable import (
    FIELD_ORIGINAL_VIDEO,
    FIELD_STATUS,
    FIELD_TITLE,
    TYPE_QUOTE,
    AirtableClient,
)
from media_publisher.sources.canva import CanvaError, parse_design_id, resolve_canva_url
from media_publisher.sources.canva_share_preview import download_canva_share_preview
from media_publisher.sources.google_drive import GoogleDriveClient
from media_publisher.sources.tn_psd import composite_without_text, safe_cache_name

FIELD_ORIGINAL_VIDEO_NAME = "Original Video Name"
FIELD_VIDEO_FOLDER = "Video Folder"
SHEET_ID = "1BGxTfnvs3zezyJVTSXroy9N0l7j5QHbzPzRj_TSjO-c"
SHEET_TAB = "English"
TN_FIELD = "pkgTn"
STATUS_KEYS = (
    "To do",
    "Translation done",
    "Editing done",
    "Synchronization done",
)
ORIGINAL_VIDEO_NAME_SUFFIX = " | Sadhguru"
FOLDER_ID_RE = re.compile(r"(?:folders/|folder/)([a-zA-Z0-9_-]+)")
WORD_DOC_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
CANVA_RE = re.compile(r"https?://(?:www\.)?(?:canva\.com|canva\.link)[^\s\"']*", re.I)
PDF_MIME = "application/pdf"
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tif", ".tiff", ".psd"}
THUMBNAIL_FILE_EXTENSIONS = IMAGE_EXTENSIONS | {".pdf"}
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "downloads" / "tn-cache"
JPG_QUALITY = 92


def normalize_title(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.endswith(ORIGINAL_VIDEO_NAME_SUFFIX):
        text = text[: -len(ORIGINAL_VIDEO_NAME_SUFFIX)].rstrip()
    return " ".join(text.casefold().split()) or None


def normalize_url(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    parsed = urlparse(text)
    if not parsed.scheme:
        return text.casefold()
    return urlunparse(
        (
            parsed.scheme.casefold(),
            parsed.netloc.casefold(),
            parsed.path.rstrip("/"),
            "",
            parsed.query,
            "",
        )
    )


def status_bucket(status: object) -> str | None:
    text = str(status or "")
    for key in STATUS_KEYS:
        if key.casefold() in text.casefold():
            return key
    return None


def build_filter_formula() -> str:
    clauses = [f'FIND("{key}", {{Status}} & "")' for key in STATUS_KEYS]
    return f"AND(OR({', '.join(clauses)}), {{Type}} != \"{TYPE_QUOTE}\")"


def tn_is_marked(value: object) -> bool:
    text = str(value or "").strip()
    return bool(text) and text.upper() != "X"


def parse_folder_id(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    match = FOLDER_ID_RE.search(text)
    if match:
        return match.group(1)
    if re.fullmatch(r"[a-zA-Z0-9_-]{10,}", text):
        return text
    return None


def is_thumbnail_file(name: str, mime_type: str) -> bool:
    if mime_type == PDF_MIME:
        return True
    if mime_type.startswith("image/"):
        return True
    if "photoshop" in mime_type.casefold():
        return True
    return Path(name).suffix.casefold() in THUMBNAIL_FILE_EXTENSIONS


def fetch_catalog_records() -> list[dict]:
    creds = service_account.Credentials.from_service_account_file(
        str(PROJECT_ROOT / "credentials" / "google-sheets-service-account.json"),
        scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"],
    )
    service = build("sheets", "v4", credentials=creds, cache_discovery=False)
    values = (
        service.spreadsheets()
        .values()
        .get(spreadsheetId=SHEET_ID, range=SHEET_TAB)
        .execute()
        .get("values", [])
    )
    return rows_to_records(values[0], values[1:]) if values else []


def index_catalog(records: list[dict]):
    by_url: dict[str, list[dict]] = defaultdict(list)
    by_title: dict[str, list[dict]] = defaultdict(list)
    for record in records:
        if not record.get("pkgSmLk"):
            continue
        url_key = normalize_url(record.get("ctLink"))
        if url_key:
            by_url[url_key].append(record)
        title_key = normalize_title(record.get("ctTitle"))
        if title_key:
            by_title[title_key].append(record)
    return by_url, by_title


def match_catalog_row(fields: dict, by_url, by_title) -> dict | None:
    for candidate in (fields.get(FIELD_ORIGINAL_VIDEO), fields.get("Original Video")):
        url_key = normalize_url(candidate)
        if url_key and by_url.get(url_key):
            return by_url[url_key][0]
    for candidate in (
        fields.get(FIELD_ORIGINAL_VIDEO_NAME),
        fields.get(FIELD_TITLE),
        fields.get("Title"),
    ):
        title_key = normalize_title(candidate)
        if title_key and by_title.get(title_key):
            return by_title[title_key][0]
    return None


def read_word_document(drive, doc) -> Document | None:
    if doc.mime_type == WORD_DOC_MIME:
        content = drive._drive.files().get_media(
            fileId=doc.id, supportsAllDrives=True
        ).execute()
        return Document(io.BytesIO(content))
    if doc.mime_type == GOOGLE_DOC_MIME:
        content = (
            drive._drive.files()
            .export_media(fileId=doc.id, mimeType=WORD_DOC_MIME)
            .execute()
        )
        return Document(io.BytesIO(content))
    return None


def paragraph_urls(paragraph: Paragraph) -> list[str]:
    urls: list[str] = []
    for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
        rel_id = hyperlink.get(qn("r:id"))
        if rel_id and rel_id in paragraph.part.rels:
            urls.append(paragraph.part.rels[rel_id].target_ref)
    urls.extend(CANVA_RE.findall(paragraph.text))
    return urls


def extract_canva_links(document: Document) -> tuple[list[str], list[str]]:
    all_urls: list[str] = []
    below_tn_urls: list[str] = []
    blocks: list[tuple[str, object]] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("p", Paragraph(child, document)))
        elif isinstance(child, CT_Tbl):
            grid = [[c.text.strip() for c in r.cells] for r in Table(child, document).rows]
            blocks.append(("t", grid))

    for i, (kind, payload) in enumerate(blocks):
        if kind == "p":
            urls = [u for u in paragraph_urls(payload) if CANVA_RE.search(u)]
            all_urls.extend(urls)
            if payload.text.strip() == "TN":
                for j in range(i + 1, len(blocks)):
                    next_kind, next_payload = blocks[j]
                    if next_kind == "t":
                        break
                    urls = [u for u in paragraph_urls(next_payload) if CANVA_RE.search(u)]
                    below_tn_urls.extend(urls)
        elif kind == "t":
            for row in payload:
                for cell in row:
                    all_urls.extend(CANVA_RE.findall(cell))

    return list(dict.fromkeys(all_urls)), list(dict.fromkeys(below_tn_urls))


def document_sort_key(item) -> tuple[int, str]:
    return (0 if item.name.upper().startswith("TEXT_") else 1, item.name.casefold())


def save_jpg(image: Image.Image, destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    rgb = image.convert("RGB")
    rgb.save(destination, format="JPEG", quality=JPG_QUALITY, optimize=True)
    return destination


def flatten_psd(path: Path) -> Image.Image:
    from psd_tools import PSDImage

    from media_publisher.sources.tn_psd import _is_artboard

    psd = PSDImage.open(path)
    artboards = [layer for layer in psd if _is_artboard(layer)]
    target = artboards[0] if artboards else psd
    return composite_without_text(target)


def flatten_pdf(path: Path) -> Image.Image:
    try:
        import fitz

        document = fitz.open(path)
        try:
            page = document[0]
            pixmap = page.get_pixmap(dpi=150)
            return Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        finally:
            document.close()
    except Exception:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(str(path))
        try:
            page = document[0]
            bitmap = page.render(scale=150 / 72)
            return bitmap.to_pil().convert("RGB")
        finally:
            document.close()


def local_source_path(tmp_path: Path, name: str, mime_type: str) -> Path:
    local = tmp_path / safe_cache_name(name)
    suffix = local.suffix.casefold()
    if suffix in THUMBNAIL_FILE_EXTENSIONS:
        return local
    if mime_type == PDF_MIME:
        return local.with_suffix(".pdf")
    if mime_type.startswith("image/"):
        return local.with_suffix(".jpg")
    if "photoshop" in mime_type.casefold():
        return local.with_suffix(".psd")
    return local


def flatten_drive_file(path: Path, *, mime_type: str = "") -> Image.Image:
    suffix = path.suffix.casefold()
    if suffix == ".psd" or "photoshop" in mime_type.casefold():
        return flatten_psd(path)
    if suffix == ".pdf" or mime_type == PDF_MIME:
        return flatten_pdf(path)
    with Image.open(path) as image:
        return image.convert("RGB")


def destination_for_title(output_dir: Path, title: str) -> Path:
    return output_dir / f"{safe_cache_name(title)}.jpg"


def pick_root_thumbnail(children) -> object | None:
    candidates = [child for child in children if is_thumbnail_file(child.name, child.mime_type)]
    if not candidates:
        return None

    def rank(child) -> tuple[int, str]:
        name = child.name.casefold()
        if name.startswith("tn_"):
            return (0, name)
        if "thumb" in name:
            return (1, name)
        if child.name.casefold().endswith(".pdf"):
            return (2, name)
        return (3, name)

    return sorted(candidates, key=rank)[0]


def canva_design_id(canva_url: str) -> str:
    resolved = resolve_canva_url(canva_url)
    design_id = parse_design_id(resolved)
    if design_id is None:
        raise CanvaError(f"Could not parse Canva design id from {canva_url!r}")
    return design_id


def main() -> int:
    settings = load_settings(PROJECT_ROOT)
    airtable = AirtableClient(
        settings.airtable_token,
        settings.airtable_base_id,
        settings.airtable_table_name,
    )
    drive = GoogleDriveClient.from_service_account(
        PROJECT_ROOT / "credentials" / "google-sheets-service-account.json"
    )
    canva_client = canva_client_from_settings(settings) if canva_settings_complete(settings) else None

    catalog = fetch_catalog_records()
    by_url, by_title = index_catalog(catalog)
    folder_cache: dict[str, list] = {}
    doc_cache: dict[str, tuple[list[str], list[str]]] = {}

    targets: list[dict] = []
    for record in airtable.list_records(filter_formula=build_filter_formula()):
        fields = record.fields
        bucket = status_bucket(fields.get(FIELD_STATUS))
        if bucket is None:
            continue
        sheet_row = match_catalog_row(fields, by_url, by_title)
        if sheet_row is None or not tn_is_marked(sheet_row.get(TN_FIELD)):
            continue
        targets.append(
            {
                "title": str(
                    fields.get(FIELD_ORIGINAL_VIDEO_NAME)
                    or fields.get(FIELD_TITLE)
                    or "Untitled"
                ).strip(),
                "status": bucket,
                "folder_id": parse_folder_id(fields.get(FIELD_VIDEO_FOLDER)),
            }
        )

    output_dir = DEFAULT_OUTPUT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict] = []
    summary = Counter()
    failures: list[tuple[str, str]] = []

    print(f"=== Cache {len(targets)} thumbnails to {output_dir} (JPEG for Airtable) ===\n")

    with tempfile.TemporaryDirectory(prefix="tn-cache-src-") as tmp:
        tmp_path = Path(tmp)
        for item in sorted(targets, key=lambda row: (row["status"], row["title"])):
            title = item["title"]
            destination = destination_for_title(output_dir, title)
            folder_id = item["folder_id"]
            source_type = "unknown"
            source_detail = ""

            try:
                if folder_id is None:
                    raise RuntimeError("missing Video Folder link")

                if folder_id not in folder_cache:
                    folder_cache[folder_id] = drive.list_children(folder_id)
                children = folder_cache[folder_id]
                root_file = pick_root_thumbnail(children)

                if root_file is not None:
                    local_source = local_source_path(
                        tmp_path, root_file.name, root_file.mime_type
                    )
                    drive.download_file(root_file.id, local_source)
                    image = flatten_drive_file(
                        local_source, mime_type=root_file.mime_type
                    )
                    save_jpg(image, destination)
                    if root_file.name.casefold().endswith(".pdf"):
                        source_type = "drive-pdf"
                    elif root_file.name.casefold().endswith(".psd"):
                        source_type = "drive-psd"
                    else:
                        source_type = "drive-image"
                    source_detail = root_file.name
                else:
                    docs = sorted(
                        [
                            child
                            for child in children
                            if child.mime_type in (WORD_DOC_MIME, GOOGLE_DOC_MIME)
                            and child.name.upper().startswith("TEXT_")
                        ],
                        key=document_sort_key,
                    )
                    if not docs:
                        raise RuntimeError("no root thumbnail file or TEXT_ doc")
                    doc = docs[0]
                    if doc.id not in doc_cache:
                        document = read_word_document(drive, doc)
                        doc_cache[doc.id] = (
                            extract_canva_links(document) if document else ([], [])
                        )
                    canva_any, canva_below_tn = doc_cache[doc.id]
                    canva_url = (canva_below_tn or canva_any or [None])[0]
                    if not canva_url:
                        raise RuntimeError("no Canva link in TEXT_ doc")
                    temp_export = tmp_path / f"{safe_cache_name(title)}.jpg"
                    source_type = "canva-share-preview"
                    if canva_client is not None:
                        try:
                            design_id = canva_design_id(canva_url)
                            canva_client.download_design_image(
                                design_id,
                                temp_export,
                                export_format="jpg",
                            )
                            source_type = "canva-export"
                        except CanvaError:
                            download_canva_share_preview(canva_url, temp_export)
                    else:
                        download_canva_share_preview(canva_url, temp_export)
                    with Image.open(temp_export) as image:
                        save_jpg(image, destination)
                    source_detail = canva_url

                summary["ok"] += 1
                print(f"OK  {title}")
                print(f"    -> {destination.name} ({source_type}: {source_detail[:100]})")
                manifest.append(
                    {
                        "title": title,
                        "status": item["status"],
                        "source_type": source_type,
                        "source_detail": source_detail,
                        "airtable_file": str(destination),
                    }
                )
            except Exception as exc:
                summary["failed"] += 1
                failures.append((title, str(exc)))
                print(f"FAIL {title}")
                print(f"     {exc}")

    manifest_path = output_dir / "pkgtn-thumbnails-manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    print()
    print("=== Summary ===")
    print(f"Cached JPGs:  {summary['ok']}")
    print(f"Failed:       {summary['failed']}")
    print(f"Manifest:     {manifest_path}")
    if failures:
        print()
        print("Failures:")
        for title, reason in failures:
            print(f"  - {title}: {reason}")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
